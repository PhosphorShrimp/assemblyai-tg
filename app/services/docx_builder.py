"""DOCX builder for transcription output."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from app.services.transcript_formatter import FormattedTranscript


def build_transcript_docx(
    output_path: Path,
    source_file_name: str,
    transcript: FormattedTranscript,
    processed_at: datetime | None = None,
) -> Path:
    """Build DOCX file with transcript metadata and content."""
    processed_at = processed_at or datetime.now()

    document = Document()
    document.add_heading("Транскрипция аудио", level=1)

    document.add_paragraph(f"Имя исходного файла: {source_file_name}")
    document.add_paragraph(f"Дата обработки: {processed_at.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Transcript ID: {transcript.transcript_id}")

    document.add_heading("Полный текст", level=2)
    document.add_paragraph(transcript.full_text or "(пусто)")

    if transcript.utterances:
        document.add_heading("Разбивка по спикерам", level=2)
        for utterance in transcript.utterances:
            document.add_paragraph(f"Спикер {utterance.speaker}: {utterance.text}")

    document.save(output_path)
    return output_path
